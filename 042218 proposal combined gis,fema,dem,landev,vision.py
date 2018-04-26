#fix save path deeds DONE ZZ
#update for multiple sections
#nereval option
#few loops would be better.
#CLEAN UP FORMATTING

#COPYRIGHT NARRAGANSETT ENGINEERING INC
#NEI-CDS.COM

#WORK PRODUCT. NOT FOR DISTRIBUTION




#load modules. make sure pip install



import time
from time import sleep
import datetime
from PIL import Image

import argparse
import re
import os
from selenium import webdriver
from selenium.common.exceptions import NoSuchElementException
from selenium.common.exceptions import NoSuchElementException
from selenium import webdriver
from selenium.webdriver.common.by import By
from selenium.webdriver.support.ui import WebDriverWait
from selenium.webdriver.support import expected_conditions as EC
from selenium.webdriver.common.keys import Keys
from selenium.webdriver.common.keys import Keys
import pathlib
import docx
import re
import os
from docx.enum.text import WD_LINE_SPACING
from docx.shared import Pt
from docx import Document
from docx.shared import Inches
import json
import datetime
from copy import deepcopy
from docx.shared import Pt
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_LINE_SPACING
import json

import datetime

now = datetime.datetime.now()

#for docx
document = Document()

#text input

print ("program searches vision apprisal in Portsmouth, Middletown, Tiverton, Newport, Narragansett; gets gis, RIDEM soils, wetlands and aerial")
print ("MUST BE FIRST HIT FOR PROGRAM TO WORK")
print  ("if using Tiverton use full name e.g. road, way, lane")
print ("if using Portsmouth Middletown or Newport use abbrev. eg. RD, WAY, LN")
print ("litteral searches for address only - not for use with vacant land - NEW FEATURES BEING ADDED. MAYBE TRY SEARCHING AND CONFIRMING ADDRESS IS FIRST HIT")

#location input

print ("program search land evidence in Portsmouth, Middletown, Tiverton, Narragansett and Newport (ALL VISION TOWNS NOW); gets lot info and screenshot of up to 5 deeds")
address = input("Enter Address")
town = input("Enter Town, no spaces, and capitalize first letter")
state = input("Enter Two Letter State")

print("open town page" + town)

#REF WORD START TO LOAD FULLTEXT

#reads docx and converts to text variable fullText
def getText(filename):
    doc = docx.Document(filename)
    fullText = []
    for para in doc.paragraphs:
        fullText.append(para.text)
    return '\n'.join(fullText)

if __name__ == '__main__':
    filename='NEI - MASTER PROPOSAL SCOPE - HOURLY RATES - NOTES 2018.docx';  #docx file name
    fullText=getText(filename)
    # print (fullText)

print ('loading master proposal list')


#proposal section input

sect1 = input("Enter Section1 - three digit code from master for proposal section - three numbers USE 0 FOR '.' delimieter (E.G. FOR sect 1.1  YOU WOULD ENTER: 101")  #first section

#additional section NOT QUITE WORKING
cont = input ('enter Y/N - would you like to add another section  - SAME FORMAT?')

if cont == "Y":
    sect2 = input ('enter sect2')
    match2 = re.findall(str(sect2) +'(.*?)' + str(sect2), str(fullText), re.DOTALL)  #adds section no.
    p2 = re.sub("\n", " ", str(''.join(match2)), re.DOTALL)
    print ('sect2')
    print (sect2)

if cont == "N":
    sect2 = "0"
    p2 = "0"
    print ('ok')
else:
    sect2 = "0"  #MAYBE DELETE THIS
    p2 = "0"
    print ('done')


pathlib.Path('c:/python36/docs/' + address + '/').mkdir(parents=True, exist_ok=True)


driver = webdriver.Firefox(executable_path=r'c:\gecko\geckodriver.exe')

if town == 'Portsmouth':
    web = 'https://i2e.uslandrecords.com/RI/Portsmouth/D/Default.aspx'
elif town == 'Middletown':
    web = 'https://i2f.uslandrecords.com/RI/Middletown/D/Default.aspx'
elif town == 'Tiverton':
    web = 'https://i2f.uslandrecords.com/RI/Tiverton/D/Default.aspx'
elif town == 'Narragansett':
    web = 'https://i2e.uslandrecords.com/RI/Narragansett/D/Default.aspx'
elif town == 'Newport':
    web = 'https://i2f.uslandrecords.com/RI/Newport'
elif town == 'Providence':
    web = 'https://i2f.uslandrecords.com/RI/Providence/D/Default.aspx'
else:
    web = "0"
    print ('you spelled the Town name wrong or did not capitalize dummy')

print ('searching land records in')
print (town)
print (web)

    



#START HERE FOR RIDEM GIS MAP ---------------------------------------------------------
#RIDEM WETLANDS - SOILS - FLOOD ZONE MAP

driver.get('http://ridemgis.maps.arcgis.com/apps/webappviewer/index.html?id=87e104c8adb449eb9f905e5f18020de5') #get webpage mapss
driver.maximize_window()
time.sleep(2)

wait = WebDriverWait(driver, 10)
wait.until(EC.presence_of_element_located((By.XPATH, '//*[@id="esri_dijit_Search_0_input"]')))  #waits address button

element = driver.find_element_by_xpath('//*[@id="esri_dijit_Search_0_input"]') #finds address
element.send_keys(address + " " + town + " " + state) #sends address
element.send_keys(Keys.RETURN) #hits enter


wait = WebDriverWait(driver, 10)
wait.until(EC.presence_of_element_located((By.XPATH, '//*[@id="jimu_dijit_CheckBox_68"]/div[1]')))  #waits for soils button

element = driver.find_element_by_xpath('//*[@id="jimu_dijit_CheckBox_68"]/div[1]') #find soils button
element.click() #click soils button


element = driver.find_element_by_xpath('//*[@id="jimu_dijit_CheckBox_130"]/div[1]') #find  2014 aerial
element.click() #click button

element = driver.find_element_by_xpath('//*[@id="jimu_dijit_CheckBox_55"]/div[1]') #find  RIEMA FLOOD ZONE
element.click() #click button

element = driver.find_element_by_xpath('//*[@id="jimu_dijit_CheckBox_87"]/div[1]') #find  WETLANDS
element.click() #click button


time.sleep(2)

element = driver.find_element_by_xpath('/html/body/div[2]/div/div[2]/div[2]/div[11]/img') #find legend    # ADD CONDITIONAL WAITS
element.click() #click button

time.sleep(4)


driver.get_screenshot_as_file('c:/python36/docs/' + address + '/' + 'ridemgis' + '.png')   #saves first screnshot

print ('done saved screenshot RIDEM')
driver.quit()

#END RIDEM WETLANDS----------------------------------

#START HERE FOR FEMA MAP--------------------------------------------

from selenium import webdriver
driver = webdriver.Firefox(executable_path=r'c:\gecko\geckodriver.exe')
driver = webdriver.Firefox()

from selenium.webdriver.common.keys import Keys
from selenium.webdriver.common.keys import Keys

driver.get('https://msc.fema.gov/portal/search')

driver.maximize_window()

elem = driver.find_element_by_id('txtfloodmapsearch')
elem.clear()
elem.send_keys(address + " " + town + " " + state)
elem.send_keys(Keys.RETURN)

driver.execute_script("window.scrollTo(0,2000);")

time.sleep(3)

driver.get_screenshot_as_file  ('fema_map.png')

driver.quit()

print ("Done FEMA")

#DONE FEMA-------------------------------------------------------

#START GIS-------------------------------------------------------



print("open town page" + town)

driver = webdriver.Firefox(executable_path=r'c:\gecko\geckodriver.exe')
driver = webdriver.Firefox()

if town == 'Portsmouth':
    gis = 'http://www.mainstreetmaps.com/ri/portsmouth/public.asp'
elif town == 'Tiverton':
    gis = 'http://www.mainstreetmaps.com/ri/tiverton/public.asp'
elif town == 'Newport':  #arcgis
    gis = 'http://newportri.maps.arcgis.com/apps/webappviewer/index.html?id=78f7978f5667474da76d2533481662e4'
elif town == 'Middletown':  #tighebond
    gis = 'http://hosting.tighebond.com/MiddletownRI_Public/index.html'
elif town == 'Warren': #mainstmaps
    gis = 'http://www.mainstreetmaps.com/ri/warren/public.asp'
elif town == 'Providence': #arcgis
    gis = 'https://pvdgis.maps.arcgis.com/apps/webappviewer/index.html?id=1700e1cf7228491d962973edc9924484'
else:
    print ('you spelled the Town name wrong or did not capitalize dummy OR town not in gis index')

driver = webdriver.Firefox()
driver.maximize_window()

#town list
mainst = ["Portsmouth", "Tiverton", "Warren"]
arcgis = ["Newport", "Providence"]
tighebond = ["Middletown"]

if town in mainst:
    print("Using MainStMaps GIS")
    driver.get(gis)
    print ("searching for")
    print (town)
    print (gis)
    element = driver.find_element_by_id('d_disc_ok')
    element.click()

    elem = driver.find_element_by_id('s_location')
    elem.clear()
    elem.send_keys(address)
    elem.send_keys(Keys.DOWN)
    elem.send_keys(Keys.RETURN)

    time.sleep(3)

    elem = driver.find_element_by_id('baselayers')
    elem.send_keys(Keys.DOWN)
    elem.send_keys(Keys.RETURN)

    time.sleep(3)

    driver.get_screenshot_as_file ('c:/python36/docs/' + address + '/' + 'tempgis.png')  #TEMP screenshot location
    print ("tempgis screenshot saved")
    
elif town in arcgis:   #arcgis search
    print("Using ARCGIS GIS")
    driver.get(gis)
    print ("searching for")
    print (town)
    print (gis)
    time.sleep(3)
    element = driver.find_element_by_xpath('//*[@id="esri_dijit_Search_0_input"]')  #click search bar
    element.click()
    element.send_keys(address)
    #elem.send_keys(Keys.UP)
    element.send_keys(Keys.RETURN)
    time.sleep(2)

    """element = driver.find_element_by_xpath('//*[@id="widgets_ZoomSlider_Widget_31"]/div[1]')  #one click zoom in
    element.click()
    element = driver.find_element_by_xpath('//*[@id="widgets_ZoomSlider_Widget_31"]/div[1]')  #one click zoom in
    element.click()
    element = driver.find_element_by_xpath('//*[@id="widgets_ZoomSlider_Widget_31"]/div[1]')  #one click zoom in
    element.click()"""  #ZOOM DOES NOT WORK IN PROVIDENCE
    time.sleep(3)
    
    driver.get_screenshot_as_file ('c:/python36/docs/' + address + '/' + 'tempgis.png')  #screenshot location
    print ("tempgis screenshot saved")

elif town in tighebond:  #tighebond search   #lot more to to click easements, etc. LOT OF INFO
    print("Using tighebond GIS")
    driver.get(gis)
    print ("searching for")
    print (town)
    print (gis)
    time.sleep(3)
    element = driver.find_element_by_xpath('//*[@id="searchinput"]')  #click search bar
    element.click()
    element.send_keys(address)
    #elem.send_keys(Keys.UP)
    element.send_keys(Keys.RETURN)
    time.sleep(3)
    element = driver.find_element_by_xpath('//*[@id="tabbasemap"]/button/div')  #click layer bar
    element.click()
    element = driver.find_element_by_xpath('//*[@id="baseMapGallery"]/li[4]/a/img')  #click googlem map bar
    element.click()
    
    
    time.sleep(3)
    driver.get_screenshot_as_file('c:/python36/docs/' + address + '/' + 'tempgis.png')
    print ("tempgis screenshot saved")

  
   
else:
    print("town not in gis list")

driver.quit()



#END GIS ---------------------------

#START VISION --------------------------------

#town list
Vision = ["Burrillville", "Charlestown", "Cranston", "Cumberland", "East Providence", "Exeter", "Foster", "Glocester", "Hopkinton", "Jamestown", "Johnston", "Little Compton", "Lincoln" 
,"Middletown", "Narragansett", "Newport", "New Shoreham", "North Kingstown", "Pawtucket", "Portsmouth",  "Richmond", "Smithfield", "South Kingstown", "Tiverton", "Warwick" ,"Westerly", "Woonsocket"]  


if town in Vision:
    print("Using VisionAppriasal Database")
    driver = webdriver.Firefox(executable_path=r'c:\gecko\geckodriver.exe')
    driver = webdriver.Firefox()

    driver.get('http://www.vgsi.com/vision/applications/parceldata/RI/Home.aspx')
    driver.maximize_window()

    element = driver.find_element_by_partial_link_text(town)
    element.click()

    element = driver.find_element_by_id('MainContent_btnEnterOnlineDatabase')
    element.click()

    elem = driver.find_element_by_id('MainContent_txtSearchAddress')


    elem.send_keys(address)
    time.sleep(2)
    elem.send_keys(Keys.UP)
    time.sleep(2)
    elem.send_keys(Keys.RETURN)


    time.sleep(3)

    ###scroll and screenshots


    driver.get_screenshot_as_file('1vision appraisal.png')

    driver.execute_script("window.scrollTo(0,1000);")


    driver.get_screenshot_as_file('2vision appraisal.png')

    driver.execute_script("window.scrollTo(0,1200);")

    time.sleep(3)

    driver.execute_script("window.scrollTo(0,2400);")

    driver.get_screenshot_as_file('3vision appraisal.png')

    #get relevant info owner, co, zone, area, plat lot
    #vars: own, co, ap, bp, zone, area


    own = driver.find_element_by_id('MainContent_lblOwner').text
    print (own)

    co = driver.find_element_by_id('MainContent_lblCoOwner').text
    print (co)

    ap = driver.find_element_by_id('MainContent_lblMblu').text
    print (ap)

    bp1 = driver.find_element_by_id('MainContent_lblBp').text
    print ("book and page_" + bp1)

    zone = driver.find_element_by_id('MainContent_lblZone').text
    print (zone)

    ac = driver.find_element_by_id('MainContent_lblLndAcres').text
    print (ac)

    #splits book page with / - as delimiters

    #BOOKPAGE1   #if element greater than 0, get element and parse by delim. if not exist print no data

    if len(driver.find_elements_by_xpath('//*[@id="MainContent_grdSales"]/tbody/tr[2]/td[4]')) > 0: 
        bp1 = driver.find_element_by_xpath('//*[@id="MainContent_grdSales"]/tbody/tr[2]/td[4]').text
        print (bp1)
    else:
        bp1 = "x"
        book1 = "x"
        page1 = "x"
        print ('no bp1')

    if len(bp1) >= 3:    
        book1, page1 = re.split("[-/]", bp1)
        print ("book_" + book1)
        print ("page_" + page1)
    else:
        print ('no data')
        
    #BOOKPAGE2
    if len(driver.find_elements_by_xpath('//*[@id="MainContent_grdSales"]/tbody/tr[3]/td[4]')) > 0: 
        bp2 = driver.find_element_by_xpath('//*[@id="MainContent_grdSales"]/tbody/tr[3]/td[4]').text
        print (bp2)
    else:
        bp2 = "x"
        book2 = "x"
        page2 = "x"
        print ('no bp2')

    if len(bp2) >= 3:    
        book2, page2 = re.split("[-/]", bp2)
        print ("book_" + book2)
        print ("page_" + page2)
    else:
        print ('no data')
        
    #BOOKPAGE3
    if len(driver.find_elements_by_xpath('//*[@id="MainContent_grdSales"]/tbody/tr[4]/td[4]')) > 0: 
        bp3 = driver.find_element_by_xpath('//*[@id="MainContent_grdSales"]/tbody/tr[4]/td[4]').text
        print (bp3)
    else:
        bp3 = "x"
        book3 = "x"
        page3 = "x"
        print ('no bp3')

    if len(bp3) >= 3:    
        book3, page3 = re.split("[-/]", bp3)
        print ("book_" + book3)
        print ("page_" + page3)
    else:
        print ('no data')



    #BOOKPAGE4
       
    if len(driver.find_elements_by_xpath('//*[@id="MainContent_grdSales"]/tbody/tr[5]/td[4]')) > 0: 
        bp4 = driver.find_element_by_xpath('//*[@id="MainContent_grdSales"]/tbody/tr[5]/td[4]').text
        print (bp4)
    else:
        bp4 = "x"
        book4 = "x"
        page4 = "x"
        print ('no bp4')

    if len(bp4) >= 3:    
        book4, page4 = re.split("[-/]", bp4)
        print ("book_" + book4)
        print ("page_" + page4)
    else:
        print ('no data')

    #BOOKPAGE5
        
    if len(driver.find_elements_by_xpath('//*[@id="MainContent_grdSales"]/tbody/tr[5]/td[4]')) > 0: 
        bp5 = driver.find_element_by_xpath('//*[@id="MainContent_grdSales"]/tbody/tr[5]/td[4]').text
        print (bp5)
    else:
        bp5 = "x"
        book5 = "x"
        page5 = "x"
        print ('no bp4')
        

    if len(bp5) >= 3:    
        book5, page5 = re.split("[-/]", bp4)
        print ("book_" + book5)
        print ("page_" + page5)
    else:
        print ('no data')

    driver.quit()

#END VISION ---------------------

if town  == "Providence":  #using regular method, can't use it, finds EAST PROVIDENCE FIRST, HENCE WORKAROUND

    print("Using VisionAppriasal Database")
    driver = webdriver.Firefox(executable_path=r'c:\gecko\geckodriver.exe')
    driver = webdriver.Firefox()

    driver.get('http://gis.vgsi.com/providenceri/')
    driver.maximize_window()

    element = driver.find_element_by_id('MainContent_btnEnterOnlineDatabase')
    element.click()

    elem = driver.find_element_by_id('MainContent_txtSearchAddress')


    elem.send_keys(address)
    time.sleep(2)
    elem.send_keys(Keys.DOWN)
    time.sleep(2)
    elem.send_keys(Keys.RETURN)


    time.sleep(3)

    ###scroll and screenshots


    driver.get_screenshot_as_file('1vision appraisal.png')

    driver.execute_script("window.scrollTo(0,1000);")


    driver.get_screenshot_as_file('2vision appraisal.png')

    driver.execute_script("window.scrollTo(0,1200);")

    time.sleep(3)

    driver.execute_script("window.scrollTo(0,2400);")

    driver.get_screenshot_as_file('3vision appraisal.png')

    #get relevant info owner, co, zone, area, plat lot
    #vars: own, co, ap, bp, zone, area


    own = driver.find_element_by_id('MainContent_lblOwner').text
    print (own)

    co = driver.find_element_by_id('MainContent_lblCoOwner').text
    print (co)

    ap = driver.find_element_by_id('MainContent_lblMblu').text
    print (ap)

    bp1 = driver.find_element_by_id('MainContent_lblBp').text
    print ("book and page_" + bp1)

    zone = driver.find_element_by_id('MainContent_lblZone').text
    print (zone)

    ac = driver.find_element_by_id('MainContent_lblLndAcres').text
    print (ac)

    #splits book page with / - as delimiters

    #BOOKPAGE1   #if element greater than 0, get element and parse by delim. if not exist print no data

    if len(driver.find_elements_by_xpath('//*[@id="MainContent_grdSales"]/tbody/tr[2]/td[4]')) > 0: 
        bp1 = driver.find_element_by_xpath('//*[@id="MainContent_grdSales"]/tbody/tr[2]/td[4]').text
        print (bp1)
    else:
        bp1 = "x"
        book1 = "x"
        page1 = "x"
        print ('no bp1')

    if len(bp1) >= 3:    
        book1, page1 = re.split("[-/]", bp1)
        print ("book_" + book1)
        print ("page_" + page1)
    else:
        print ('no data')
        
    #BOOKPAGE2
    if len(driver.find_elements_by_xpath('//*[@id="MainContent_grdSales"]/tbody/tr[3]/td[4]')) > 0: 
        bp2 = driver.find_element_by_xpath('//*[@id="MainContent_grdSales"]/tbody/tr[3]/td[4]').text
        print (bp2)
    else:
        bp2 = "x"
        book2 = "x"
        page2 = "x"
        print ('no bp2')

    if len(bp2) >= 3:    
        book2, page2 = re.split("[-/]", bp2)
        print ("book_" + book2)
        print ("page_" + page2)
    else:
        print ('no data')
        
    #BOOKPAGE3
    if len(driver.find_elements_by_xpath('//*[@id="MainContent_grdSales"]/tbody/tr[4]/td[4]')) > 0: 
        bp3 = driver.find_element_by_xpath('//*[@id="MainContent_grdSales"]/tbody/tr[4]/td[4]').text
        print (bp3)
    else:
        bp3 = "x"
        book3 = "x"
        page3 = "x"
        print ('no bp3')

    if len(bp3) >= 3:    
        book3, page3 = re.split("[-/]", bp3)
        print ("book_" + book3)
        print ("page_" + page3)
    else:
        print ('no data')



    #BOOKPAGE4
       
    if len(driver.find_elements_by_xpath('//*[@id="MainContent_grdSales"]/tbody/tr[5]/td[4]')) > 0: 
        bp4 = driver.find_element_by_xpath('//*[@id="MainContent_grdSales"]/tbody/tr[5]/td[4]').text
        print (bp4)
    else:
        bp4 = "x"
        book4 = "x"
        page4 = "x"
        print ('no bp4')

    if len(bp4) >= 3:    
        book4, page4 = re.split("[-/]", bp4)
        print ("book_" + book4)
        print ("page_" + page4)
    else:
        print ('no data')

    #BOOKPAGE5
        
    if len(driver.find_elements_by_xpath('//*[@id="MainContent_grdSales"]/tbody/tr[5]/td[4]')) > 0: 
        bp5 = driver.find_element_by_xpath('//*[@id="MainContent_grdSales"]/tbody/tr[5]/td[4]').text
        print (bp5)
    else:
        bp5 = "x"
        book5 = "x"
        page5 = "x"
        print ('no bp4')
        

    if len(bp5) >= 3:    
        book5, page5 = re.split("[-/]", bp4)
        print ("book_" + book5)
        print ("page_" + page5)
    else:
        print ('no data')

    driver.quit()

    #END PROVIDENCE WORKAROUND


else:
    print("town not in assessors list")


#ADD NE REVAL -----------------------------------------------------------------------------------------




#IMAGE CROPPING

#FEMA CROP

from PIL import Image
 
""" crops image from fema box = (left, top, left+width, top+height)"""

# size is width/height
img = Image.open('fema_map.png')
box = (450, 0, 1500, 900)
area = img.crop(box)

area.save ('c:/python36/docs/' + address + '/' + 'cropfema.png')  #CROPPED LOCATION

#VISION CROP1

""" crops image from  box = (left, top, left+width, top+height)"""

# size is width/height
img = Image.open('1vision appraisal.png')
box = (450, 0, 1440, 978)
area = img.crop(box)

area.save ('c:/python36/docs/' + address + '/' + '1VAcrop.png')

#VISION CROP2

""" crops image from  box = (left, top, left+width, top+height)"""

# size is width/height
img = Image.open('2vision appraisal.png')
box = (450, 0, 1440, 978)
area = img.crop(box)

area.save ('c:/python36/docs/' + address + '/' + '2VAcrop.png')

#VISION CROP3

""" crops image from  box = (left, top, left+width, top+height)"""

# size is width/height
img = Image.open('3vision appraisal.png')
box = (450, 0, 1440, 978)
area = img.crop(box)

area.save ('c:/python36/docs/' + address + '/' + '3VAcrop.png')



#END CROPPING

#START HERE FOR WORD DOC CREATION



#WORD DOCUMENT START ---------------------------------------------------------




#NEW 04.21.18 - INSERT SECTIONS FROM WORD DOC: ---------------------------------




###LOADS WORD DOCUMENT FROM START - REF WORD START

##############################  going to want to loop this. can only do two sections FTM
    #need to add notes to match section 3 and 4
    
#FIRST SECTION
match1 = re.findall(sect1 +'(.*?)' + sect1, str(fullText), re.DOTALL)  #adds section no.
print (match1)

p1 = re.sub("\n", " ", str(''.join(match1)), re.DOTALL)  #removes newline section
print (p1)


#STANDARD NOTES
sect0 = "s.00"   #standard notes
match0 = re.findall(sect0 +'(.*?)' + sect0, str(fullText), re.DOTALL)  #adds RAW standard notes.
  

p0 = re.sub("\n", " ", str(''.join(match0)), re.DOTALL)  #removes newline section  ADD PARSED STD NOTES
print (p0)

#SECTION NOTES

matchlist1 = ["101", "102", "103", "104", "105", "106", "107", "108"]  #MATCHLIST NO SECTION 1
matchlist2 = ["201", "202", "203", "204", "205" ,"206" , "207", "208"]  #MATCHLIST NO SECTION 2
matchlist3 = ["301", "302", "303" ,"304" ,"305" ,"306" ,"307"]

if sect1 in matchlist1: #if section 1, then section 1 std notes
    sectN1 = "s.012"
    print (sectN1)
elif sect1 in  matchlist2: #if section 1, then section 1 std notes
    sectN1 = "s.012"
    print (sectN1)
elif sect1 in  matchlist3: #if section 1, then section 1 std notes
    sectN1 = "s.03"
    print (sectN1)
else:
    print ('not in list USER ERROR')

    
match3 = re.findall(sectN1 +'(.*?)' + sectN1, str(fullText), re.DOTALL)  #adds section no. PER FIRST SECTION
print (match3)

p3 = re.sub("\n", " ", str(''.join(match3)), re.DOTALL)  #removes Newline /n for PARSED NOTES p3
print (p3)







#CREATE NEW DOCUMENT

style = document.styles['Normal']
font = style.font
font.name = 'Swis721 Lt BT'
font.size = Pt(10)


paragraph = document.add_paragraph()
paragraph.line_spacing_rule = WD_LINE_SPACING.SINGLE

p = document.add_paragraph()
r = p.add_run()
r.add_picture('c:/python36/NEI Logo-Proposal.jpg', width=Inches(5.0))





document.add_paragraph('                                             ')

p = document.add_paragraph ('Date + Time: ' + now.strftime("%Y-%m-%d %H:%M"))

document.add_paragraph('In care of:__________________________________________')
document.add_paragraph('Owner(s) of record: ' + own + " " + co + " " + address + " " + town + " " + state)
document.add_paragraph('Plat + Lot (A.P.).: ' + ap)
document.add_paragraph('                                             ')
document.add_paragraph('Site Information: ' + bp1)
document.add_paragraph('Latest Book and Page: ' + bp1)
document.add_paragraph('Land evidence chain: ' + bp2 + " " + bp3 + " " + bp4 + " " + bp5)
document.add_paragraph('Zone: ' + zone)
document.add_paragraph('Lot Area (Acres): ' + ac + '+/-')
document.add_paragraph('All information from Assessors Database unless noted')
document.add_paragraph('                                             ')
document.add_paragraph('                                             ')
document.add_paragraph('Narragansett Engineering Inc is please to provide you with the following proposal, regarding: ')

p = document.add_paragraph(town + ' GIS + Aerial Map: ')
r = p.add_run()
r.add_picture('c:/python36/docs/' + address + '/' + 'tempgis.png', width=Inches(5.0))

###START VARIABLE SECTION FROM MASTER PROPOSAL

p = document.add_paragraph(p1)  #adds new SECTION SCOPE from MASTER DOCUMENT
r = p.add_run()

if len(sect2) >= 3:
    p = document.add_paragraph(p2)  #adds SECOND SECTION - WITHOUT NOTES #_____-------------------------- ADD SECTION NOTES
    r = p.add_run()
    print ('adding second section - MAKE SURE TO ADD NOTES IF SECOND SECTION')
else:
    print ('nothing to add')

p = document.add_paragraph(p0)  #adds standard notes
r = p.add_run()

p = document.add_paragraph(p3)  #adds section notes
r = p.add_run()

###END VARIABLE SECTION FROM MASTER PROPOSAL




p = document.add_paragraph('Assesors Database Information Continued: ')
r = p.add_run()
r.add_text('From assessors datatabase: ' + now.strftime("%Y-%m-%d %H:%M"))
r.add_picture('c:/python36/docs/' + address + '/' + '1VAcrop.png', width=Inches(4.0))

p = document.add_paragraph('Assesors Database Information Continued: ')
r = p.add_run()
r.add_picture('c:/python36/docs/' + address + '/' + '2VAcrop.png', width=Inches(4.0))

p = document.add_paragraph('Assesors Database Information Continued: ')
r = p.add_run()
r.add_picture('c:/python36/docs/' + address + '/' + '3VAcrop.png', width=Inches(4.0))


p = document.add_paragraph('RIDEM GIS Soils, Wetlands, Flood Zone: ')
r = p.add_run()
r.add_picture('c:/python36/docs/' + address + '/' + 'ridemgis.png', width=Inches(4.0))

p = document.add_paragraph('FEMA GIS Information: ')
r = p.add_run()
r.add_picture('c:/python36/docs/' + address + '/' + 'cropfema.png', width=Inches(4.0))


document.add_page_break()



document.save('c:/python36/docs/' + address + '/' + 'NEI SAMPLE.docx')

print ("Done! Check NEI Sample.docx")
print ("saved in: " 'c:/python36/docs/' + address + '/')



#WORD DOCUMENT END ---------------------------------------------------------

#START COPY LOOP 1 -----------------   SEARCHES LAND EVIDENCE


#UPDATE SAVE PATH TO //LAND EVIDENCE IN DIR-----------------------------------------------------------------------------------------------------------------------

print ('back to  main window')
###add in if len(web) >1  .... else do nothing

if bp1 != "x": 

    driver = webdriver.Firefox(executable_path=r'c:\gecko\geckodriver.exe')
    driver.get(web)

    time.sleep(4)

    element = driver.find_element_by_xpath('//*[@id="SearchCriteriaName1_DDL_SearchName"]/option[3]') #set to search volume
    element.click()

    time.sleep(3)

    elem =  driver.find_element_by_id('SearchFormEx1_ACSTextBox_Volume')  #find box for book
    elem.send_keys(book1) #book1                                             #XXX

    elem = driver.find_element_by_xpath('//*[@id="SearchFormEx1_ACSTextBox_PageNumber"]') # find box for page xx
    elem.send_keys(page1)   #page1                                                          #XXX
    elem.send_keys(Keys.RETURN)

    time.sleep(6)

    try:
        element = driver.find_element_by_css_selector('#DocList1_GridView_Document_ctl02_ImgBut')   #clicks image button for deed
        element.click()

        time.sleep(6)
          
        # Get windows list and put focus on new window (which is on the 1st index in the list)
        """windows = driver.window_handles  #ERROR switching to window--------------------------------------------------
        driver.switch_to.window(windows[1])
        driver.maximize_window()"""
        
        WebDriverWait(driver, 20).until(EC.number_of_windows_to_be(2)) #wait for two windows

        newWindow = driver.window_handles
        newNewWindow = newWindow[1]
        driver.switch_to.window(newNewWindow)

        driver.maximize_window()

        print ('switching to popup')

        element = driver.find_element_by_id('ImageViewer1_BtnFitToHeight') # fits image to height
        element.click()

        time.sleep(4)  #sleep 4 sec

        driver.get_screenshot_as_file('c:/python36/docs/' + address + '/' + bp1 + '_deed.png')  #cdeed round 1 xx --------------------   ZZ
        print ("screenshot_" + bp1)                         #  XXX

        time.sleep(5)

        opt = driver.find_element_by_id('ImageViewer1_BtnNext')
        counter = 0
        while counter < 30:
            try:
                counter = counter + 1
                element = driver.find_element_by_id('ImageViewer1_BtnNext')  #click  next
                element.click()
                element = driver.find_element_by_id('ImageViewer1_BtnFitToHeight') #fit to window height
                element.click()
                
                time.sleep(4) #sleep - load for screenshot
                
                file_name = 'c:/python36/docs/' + address + '/' +  bp1 + "_" + str(counter) + "_deed" + ".png"   #save as filename xx  #UPDATED XXXX   ZZ
                driver.get_screenshot_as_file(file_name)
                print ('saving' + (file_name))

            
            except NoSuchElementException :
                pass
                break
                print ('in loop, getting next page') 
                
            else:
                print ('another try')
        else:
            print ('done with this deed')
                
        print ('now what')

    except NoSuchElementException:
        print (" ONLY ONE PAGE OR ERROR IN DATABASE or DUPLICATE ENTRIES OR MISSING IMAGE, GO FIND " + book1 + " " + page1) #XX
        
    print  ("done")
    driver.quit()

else:
    print ("no value for bp2")
    driver.quit()
    

#END COPY LOOP 1  ---------------


#START COPY LOOP 2 -----------------   

print ('back to  main window')

if bp2 != "x": 

    driver = webdriver.Firefox(executable_path=r'c:\gecko\geckodriver.exe')
    driver.get(web)

    time.sleep(4)

    element = driver.find_element_by_xpath('//*[@id="SearchCriteriaName1_DDL_SearchName"]/option[3]') #set to search volume
    element.click()

    time.sleep(3)

    elem =  driver.find_element_by_id('SearchFormEx1_ACSTextBox_Volume')  #find box for book
    elem.send_keys(book2) #book2                                             #XXX

    elem = driver.find_element_by_xpath('//*[@id="SearchFormEx1_ACSTextBox_PageNumber"]') # find box for page xx
    elem.send_keys(page2)   #page2                                                          #XXX
    elem.send_keys(Keys.RETURN)

    time.sleep(6)

    try:
        element = driver.find_element_by_css_selector('#DocList1_GridView_Document_ctl02_ImgBut')   #clicks image button for deed
        element.click()

        time.sleep(6)
          
        # Get windows list and put focus on new window (which is on the 1st index in the list)
        """windows = driver.window_handles  #ERROR switching to window--------------------------------------------------
        driver.switch_to.window(windows[1])
        driver.maximize_window()"""
        
        WebDriverWait(driver, 20).until(EC.number_of_windows_to_be(2)) #wait for two windows

        newWindow = driver.window_handles
        newNewWindow = newWindow[1]
        driver.switch_to.window(newNewWindow)

        driver.maximize_window()

        print ('switching to popup')

        element = driver.find_element_by_id('ImageViewer1_BtnFitToHeight') # fits image to height
        element.click()

        time.sleep(4)  #sleep 4 sec

        driver.get_screenshot_as_file('c:/python36/docs/' + address + '/' + bp2 + '_deed.png')  #cdeed round 2 xx -----------------------------------------------------  ZZ
        print ("screenshot_" + bp2)                         #  XXX

        time.sleep(5)

        opt = driver.find_element_by_id('ImageViewer1_BtnNext')
        counter = 0
        while counter < 30:
            try:
                counter = counter + 1
                element = driver.find_element_by_id('ImageViewer1_BtnNext')  #click  next
                element.click()
                element = driver.find_element_by_id('ImageViewer1_BtnFitToHeight') #fit to window height
                element.click()
                
                time.sleep(4) #sleep - load for screenshot
                
                file_name = 'c:/python36/docs/' + address + '/' +  bp2 + "_" + str(counter) + "_deed" + ".png"   #save as filename xx   XXXXXXXXXXXX -----------------------  ZZ
                driver.get_screenshot_as_file(file_name)
                print ('saving' + (file_name))

            
            except NoSuchElementException :
                pass
                break
                print ('in loop, getting next page') 
                
            else:
                print ('another try')
        else:
            print ('done with this deed')
                
        print ('now what')

    except NoSuchElementException:
        print (" ONLY ONE PAGE OR ERROR IN DATABASE or DUPLICATE ENTRIES OR MISSING IMAGE, GO FIND " + book2 + " " + page2) #XX
        
    print  ("done")
    driver.quit()

else:
    print ("no value for bp2")
    driver.quit()
    

#END COPY LOOP 2  ---------------


#START COPY LOOP 3 -----------------   4 first to test window


print ('back to  main window')

if bp3 != "x": 

    driver = webdriver.Firefox(executable_path=r'c:\gecko\geckodriver.exe')
    driver.get(web)

    time.sleep(3)

    element = driver.find_element_by_xpath('//*[@id="SearchCriteriaName1_DDL_SearchName"]/option[3]') #set to search volume
    element.click()

    time.sleep(3)

    elem =  driver.find_element_by_id('SearchFormEx1_ACSTextBox_Volume')  #find box for book
    elem.send_keys(book3) #book3

    elem = driver.find_element_by_xpath('//*[@id="SearchFormEx1_ACSTextBox_PageNumber"]') # find box for page
    elem.send_keys(page3)   #page3
    elem.send_keys(Keys.RETURN)

    time.sleep(5)

    try:
        element = driver.find_element_by_css_selector('#DocList1_GridView_Document_ctl02_ImgBut')   #clicks image button for deed
        element.click()

        time.sleep(3)
            
        # Get windows list and put focus on new window (which is on the 1st index in the list)     
        WebDriverWait(driver, 20).until(EC.number_of_windows_to_be(2)) #wait for two windows

        newWindow = driver.window_handles
        newNewWindow = newWindow[1]
        driver.switch_to.window(newNewWindow)

        driver.maximize_window()

        print ('switching to popup')

        element = driver.find_element_by_id('ImageViewer1_BtnFitToHeight') # fits image to height
        element.click()

        time.sleep(4)  #sleep 4 sec

        driver.get_screenshot_as_file('c:/python36/docs/' + address + '/' + bp3 + '_deed.png')  #cdeed round 3 -----------------------------------------   ZZ
        print ("screenshot_" + bp3) #save bp3

        time.sleep(5)


        opt = driver.find_element_by_id('ImageViewer1_BtnNext')
        counter = 0
        while counter < 30:
            try:
                counter = counter + 1
                element = driver.find_element_by_id('ImageViewer1_BtnNext')  #click  next
                element.click()
                element = driver.find_element_by_id('ImageViewer1_BtnFitToHeight') #fit to window height
                element.click()
                
                time.sleep(4) #sleep - load for screenshot
                
                file_name = 'c:/python36/docs/' + address + '/' + "book3_" + "page3_" + str(counter) + "_deed" + ".png"   #save as filename xx  ---------------XXXXXXXXXXXXXXXXXX  --ZZ
                driver.get_screenshot_as_file(file_name)
                print ('saving' + (file_name))

            
            except NoSuchElementException :
                pass
                break
                print ('in loop, getting next page') 
                
            else:
                print ('another try')
        else:
            print ('done with this deed')
                
        print ('now what')

    except NoSuchElementException:
        print (" ONLY ONE PAGE OR ERROR IN DATABASE or DUPLICATE ENTRIES OR MISSING IMAGE, GO FIND " + book3 + " " + page3)  ##XX
        
    print  ("done item 3")
    driver.quit()
    
else:
    print ("no value for bp3")
    driver.quit()

#END COPY LOOP 3  ---------------


#START COPY LOOP 4 -----------------   4 first to test window

print ('back to  main window')


if bp4 != "x": 

    driver = webdriver.Firefox(executable_path=r'c:\gecko\geckodriver.exe')
    driver.get(web)

    time.sleep(4)

    element = driver.find_element_by_xpath('//*[@id="SearchCriteriaName1_DDL_SearchName"]/option[3]') #set to search volume
    element.click()

    time.sleep(3)

    elem =  driver.find_element_by_id('SearchFormEx1_ACSTextBox_Volume')  #find box for book
    elem.send_keys(book4) #book4

    elem = driver.find_element_by_xpath('//*[@id="SearchFormEx1_ACSTextBox_PageNumber"]') # find box for page
    elem.send_keys(page4)   #page4
    elem.send_keys(Keys.RETURN)

    time.sleep(6)

    try:
        element = driver.find_element_by_css_selector('#DocList1_GridView_Document_ctl02_ImgBut')   #clicks image button for deed
        element.click()

        time.sleep(6)
        
        #ROUND 4  'c' -----
        # Get windows list and put focus on new window (which is on the 1st index in the list)
        """windows = driver.window_handles  #ERROR switching to window--------------------------------------------------
        driver.switch_to.window(windows[1])
        driver.maximize_window()"""
        
        WebDriverWait(driver, 20).until(EC.number_of_windows_to_be(2)) #wait for two windows

        newWindow = driver.window_handles
        newNewWindow = newWindow[1]
        driver.switch_to.window(newNewWindow)

        driver.maximize_window()

        print ('switching to popup')

        element = driver.find_element_by_id('ImageViewer1_BtnFitToHeight') # fits image to height
        element.click()

        time.sleep(4)  #sleep 4 sec

        driver.get_screenshot_as_file('c:/python36/docs/' + address + '/' + bp4 + 'deed.png')  #cdeed round 4  -----------------------------------------------XXXXXXXXXXXXXXXXXXXXXX   -----  ZZ
        print ("screenshot_" + bp4)

        time.sleep(5)


        opt = driver.find_element_by_id('ImageViewer1_BtnNext')
        counter = 0
        while counter < 30:
            try:
                counter = counter + 1
                element = driver.find_element_by_id('ImageViewer1_BtnNext')  #click  next
                element.click()
                element = driver.find_element_by_id('ImageViewer1_BtnFitToHeight') #fit to window height
                element.click()
                
                time.sleep(4) #sleep - load for screenshot
                
                file_name =  'c:/python36/docs/' + address + '/' + bp4  + str(counter) + "_deed" + ".png"   #save as filename ---------------------------------------------  ZZ
                driver.get_screenshot_as_file(file_name)
                print ('saving' + (file_name))

            
            except NoSuchElementException :
                pass
                break
                print ('in loop, getting next page') 
                
            else:
                print ('another try')
        else:
            print ('done with this deed')
                
        print ('now what')

    except NoSuchElementException:
        print (" ONLY ONE PAGE OR ERROR IN DATABASE or DUPLICATE ENTRIES OR MISSING IMAGE, GO FIND " + book4 + " " + page4)
        
    print  ("done")
    driver.quit()

else:
    print ("no value for bp4")
    driver.quit()

#END COPY LOOP 4  ---------------

    

#START COPY LOOP 5 -----------------

if bp5 != "x":     

    print ('back to  main window')

    driver = webdriver.Firefox(executable_path=r'c:\gecko\geckodriver.exe')
    driver.get(web)

    time.sleep(4)

    element = driver.find_element_by_xpath('//*[@id="SearchCriteriaName1_DDL_SearchName"]/option[3]') #set to search volume
    element.click()

    time.sleep(3)

    elem =  driver.find_element_by_id('SearchFormEx1_ACSTextBox_Volume')  #find box for book
    elem.send_keys(book5) #book5                                             #XXX

    elem = driver.find_element_by_xpath('//*[@id="SearchFormEx1_ACSTextBox_PageNumber"]') # find box for page xx
    elem.send_keys(page5)   #page5                                                          #XXX
    elem.send_keys(Keys.RETURN)

    time.sleep(6)

    try:
        element = driver.find_element_by_css_selector('#DocList1_GridView_Document_ctl02_ImgBut')   #clicks image button for deed
        element.click()

        time.sleep(6)
          
        # Get windows list and put focus on new window (which is on the 1st index in the list)
        """windows = driver.window_handles  #ERROR switching to window--------------------------------------------------
        driver.switch_to.window(windows[1])
        driver.maximize_window()"""
        
        WebDriverWait(driver, 20).until(EC.number_of_windows_to_be(2)) #wait for two windows

        newWindow = driver.window_handles
        newNewWindow = newWindow[1]
        driver.switch_to.window(newNewWindow)

        driver.maximize_window()

        print ('switching to popup')

        element = driver.find_element_by_id('ImageViewer1_BtnFitToHeight') # fits image to height
        element.click()

        time.sleep(4)  #sleep 4 sec

        driver.get_screenshot_as_file('c:/python36/docs/' + address + '/' + bp5 + '_deed.png')  #cdeed round 5 xx --------------------------------  ZZ
        print ("screenshot_" + bp5)                         #  XXX

        time.sleep(5)

        opt = driver.find_element_by_id('ImageViewer1_BtnNext')
        counter = 0
        while counter < 30:
            try:
                counter = counter + 1
                element = driver.find_element_by_id('ImageViewer1_BtnNext')  #click  next
                element.click()
                element = driver.find_element_by_id('ImageViewer1_BtnFitToHeight') #fit to window height
                element.click()
                
                time.sleep(4) #sleep - load for screenshot
                
                file_name =  'c:/python36/docs/' + address + '/' + bp5 + str(counter) + "_deed" + ".png"   #save as filename xx -------------------------------------------------------  ZZ
                driver.get_screenshot_as_file(file_name)
                print ('saving' + (file_name))

            
            except NoSuchElementException :
                pass
                break
                print ('in loop, getting next page') 
                
            else:
                print ('another try')
        else:
            print ('done with this deed')
                
        print ('now what')

    except NoSuchElementException:
        print (" ONLY ONE PAGE OR ERROR IN DATABASE or DUPLICATE ENTRIES OR MISSING IMAGE, GO FIND " + book5 + " " + page5) #XX
        
    print  ("done")
    driver.quit()

else:
    print ("no value for bp5")
    driver.quit()
    

#END COPY LOOP 5  ---------------

print ("DONE!! - wow, neal is pretty great")
driver.quit()












#try timeout exception...
"""from selenium.common.exceptions import TimeoutException
from selenium.webdriver.support.ui import WebDriverWait

try:
    element = WebDriverWait(driver, 10).until(lambda driver: driver.find_element_by_id('user_first_name'))
    # do smth with the found element
except TimeoutException:
    print "Element Not Found"
    driver.close()"""



"""
browser = webdriver.Firefox()
browser.get('https://www.google.com?q=python#q=python')
first_result = ui.WebDriverWait(browser, 15).until(lambda browser: browser.find_element_by_class_name('rc'))
first_link = first_result.find_element_by_tag_name('a')

# Save the window opener (current window)
main_window = browser.current_window_handle

# Open the link in a new window by sending key strokes on the element
first_link.send_keys(Keys.SHIFT + Keys.RETURN)

# Get windows list and put focus on new window (which is on the 1st index in the list)
windows = browser.window_handles
browser.switch_to.window(windows[1])

# do whatever you have to do on this page, we will just got to sleep for now
sleep(2)

# Close current window
browser.find_element_by_tag_name('body').send_keys(Keys.CONTROL + 'w')

# Put focus back on main window
browser.switch_to.window(main_window)"""

"""# Opens a new tab
self.driver.execute_script("window.open()")

# Switch to the newly opened tab
self.driver.switch_to.window(self.driver.window_handles[1])

# Navigate to new URL in new tab
self.driver.get("https://google.com")
# Run other commands in the new tab here
You're then able to close the original tab as follows

# Switch to original tab
self.driver.switch_to.window(self.driver.window_handles[0])

# Close original tab
self.driver.close()

# Switch back to newly opened tab, which is now in position 0
self.driver.switch_to.window(self.driver.window_handles[0])
Or close the newly opened tab

# Close current tab
self.driver.close()

# Switch back to original tab
self.driver.switch_to.window(self.driver.window_handles[0])"""





"""
#kill firefox and gecko
if (browser == "FIREFOX")) {
    try {
        Runtime.getRuntime().exec("taskkill /F /IM geckodriver.exe");
        Runtime.getRuntime().exec("taskkill /F /IM plugin-container.exe");
        Runtime.getRuntime().exec("taskkill /F /IM firefox.exe");
    } catch (IOException e) {
        e.printStackTrace();
    }
} else {
    driver.quit();
}
driver.quit()"""

